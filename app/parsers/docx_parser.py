"""
DOCX text extraction for benefit documents.
"""
from __future__ import annotations

from typing import Iterable


def _iter_paragraph_text(paragraphs: Iterable) -> list[str]:
    texts: list[str] = []
    for p in paragraphs:
        if getattr(p, "text", None):
            t = p.text.strip()
            if t:
                texts.append(t)
    return texts


def extract_text_from_docx_bytes(data: bytes) -> str:
    """
    Extract text from a DOCX file given as bytes.
    This pulls visible paragraph text; tables often appear as paragraphs.
    """
    import io
    from docx import Document  # type: ignore[import]

    doc = Document(io.BytesIO(data))
    parts: list[str] = []

    # Body paragraphs
    parts.extend(_iter_paragraph_text(doc.paragraphs))

    # Headers / footers can contain important plan info
    for section in doc.sections:
        header = getattr(section, "header", None)
        if header is not None:
            parts.extend(_iter_paragraph_text(header.paragraphs))
        footer = getattr(section, "footer", None)
        if footer is not None:
            parts.extend(_iter_paragraph_text(footer.paragraphs))

    return "\n\n".join(parts).strip()

